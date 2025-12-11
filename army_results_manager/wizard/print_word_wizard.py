# -*- coding: utf-8 -*-
from odoo import models, fields, api
from odoo.modules.module import get_module_resource
from io import BytesIO
import base64
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from odoo.exceptions import UserError
from copy import deepcopy


class PrintWordWizard(models.TransientModel):
    _name = "print.word.wizard"
    _description = "Wizard chọn mẫu in Word/Excel"

    type_action = fields.Selection([
        ('print', 'In báo cáo'),
        ('send_report', 'Gửi báo cáo')]
        , default='print', string="Chức năng", required=True
    )
    mau_in = fields.Selection(
        [('template1', 'Phụ lục 1'),
         ('template2', 'Phụ lục 2'),
         ('template3', 'Phụ lục 3'),
         ('template4', 'Phụ lục 4'),
         ('template5', 'Phụ lục 5')]
    )
    report_type = fields.Selection([
        ('week', 'Theo tuần'),
        ('month', 'Theo tháng'),
        ('year', 'Theo năm'),
    ], string="Loại báo cáo", default='week')

    year = fields.Char(string="Năm", default=lambda self: date.today().year)
    month = fields.Selection([
        ('1', 'Tháng 1'), ('2', 'Tháng 2'), ('3', 'Tháng 3'),
        ('4', 'Tháng 4'), ('5', 'Tháng 5'), ('6', 'Tháng 6'),
        ('7', 'Tháng 7'), ('8', 'Tháng 8'), ('9', 'Tháng 9'),
        ('10', 'Tháng 10'), ('11', 'Tháng 11'), ('12', 'Tháng 12'),
    ], string="Tháng")

    week = fields.Selection([
        ('1', 'Tuần 1'), ('2', 'Tuần 2'),
        ('3', 'Tuần 3'), ('4', 'Tuần 4'), ('5', 'Tuần 5'),
    ], string="Tuần")

    approver_id = fields.Many2one('hr.employee', string='Cán bộ phê duyệt',
                                  domain=[('role', '=', 'commanding_officer')])
    attachment_ids = fields.Many2many(
        'ir.attachment',
        string='Tài liệu PDF',
        domain=[('mimetype', '=', 'application/pdf')]
    )

    # ==================== Helper Functions ====================

    def action_send_report(self):
        self.ensure_one()

        if not self.approver_id:
            raise UserError("Bạn phải điền Cán bộ Phê duyệt trước khi Gửi báo !")

        if not self.attachment_ids:
            raise UserError("Bạn phải chọn ít nhất 1 file!")

        # Lấy model cha (record gốc)
        active_model = self.env.context.get('active_model')

        for att in self.attachment_ids:
            self.env['ir.attachment'].create({
                'name': att.name,
                'datas': att.datas,
                'res_model': active_model,
                'type': att.type,
                'public': True,
                'mimetype': att.mimetype,
                'approver_id': self.approver_id.id,
            })
        return {'type': 'ir.actions.client', 'tag': 'soft_reload'}

    @api.onchange('report_type')
    def _onchange_report_type(self):
        if self.report_type:
            self.week = self.month = False

    def replace_placeholder_with_text(self, doc, placeholder, replacement_text):
        """Thay thế placeholder trong cả paragraphs và tables, xử lý trường hợp placeholder bị split"""
        found = False

        def replace_in_paragraph(paragraph):
            """Helper function để thay thế trong một paragraph"""
            nonlocal found

            # Ghép tất cả runs lại để tìm placeholder
            full_text = ''.join(run.text for run in paragraph.runs)

            # Kiểm tra có chứa placeholder không
            if placeholder in full_text:
                found = True

                # Thay thế placeholder
                new_text = full_text.replace(placeholder, str(replacement_text))

                if paragraph.runs:
                    # Lưu format của run đầu tiên (hoặc run có format chính)
                    first_run = paragraph.runs[0]

                    saved_format = {
                        'name': first_run.font.name,
                        'size': first_run.font.size,
                        'bold': first_run.font.bold,
                        'italic': first_run.font.italic,
                        'underline': first_run.font.underline,
                    }

                    # Lưu màu chữ (có thể None)
                    try:
                        if first_run.font.color and first_run.font.color.rgb:
                            saved_format['color'] = first_run.font.color.rgb
                        else:
                            saved_format['color'] = None
                    except:
                        saved_format['color'] = None

                    # Xóa tất cả runs hiện tại
                    while len(paragraph.runs) > 0:
                        paragraph._element.remove(paragraph.runs[0]._element)

                    # Tạo run mới với text đã thay thế
                    new_run = paragraph.add_run(new_text)

                    # Khôi phục format
                    if saved_format['name']:
                        new_run.font.name = saved_format['name']
                    if saved_format['size']:
                        new_run.font.size = saved_format['size']
                    new_run.font.bold = saved_format['bold']
                    new_run.font.italic = saved_format['italic']
                    new_run.font.underline = saved_format['underline']
                    if saved_format['color']:
                        new_run.font.color.rgb = saved_format['color']

        # Thay thế trong tất cả paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)

        # Thay thế trong tất cả tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)

        return found

    def print_table(self, doc, table_index):
        """
        In ra thông tin của table

        Args:
            doc: Document object
            table_index: Vị trí table (0-based, table_index=1 là table thứ 2)
        """
        if table_index >= len(doc.tables):
            print(f"Table index {table_index} không tồn tại!")
            print(f"Document chỉ có {len(doc.tables)} tables")
            return False

        table = doc.tables[table_index]

        print("=" * 80)
        print(f"TABLE INDEX: {table_index}")
        print(f"Số dòng: {len(table.rows)}")
        print(f"Số cột: {len(table.columns)}")
        print("=" * 80)

        # In ra từng dòng và cell
        for row_idx, row in enumerate(table.rows):
            print(f"\n--- Dòng {row_idx} ---")
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                print(f"  Cell[{row_idx}][{col_idx}]: {cell_text}")

        print("=" * 80)
        return True

    # ==================== Main Action ====================

    def action_print_word(self):
        if self.type_action == 'print' and not self.report_type:
            raise UserError('Bạn phải chọn Loại báo cáo trước khi in!')

        # Báo cáo năm
        if self.report_type == 'year' and not self.year:
            raise UserError('Bạn phải điền năm trước khi in Báo cáo năm!')

        # Báo cáo tháng
        if self.report_type == 'month' and (not self.year or not self.month):
            raise UserError('Bạn phải điền năm và tháng trước khi in Báo cáo theo tháng!')

        # Báo cáo tuần
        if self.report_type == 'week' and (not self.year or not self.month or not self.week):
            raise UserError('Bạn phải điền năm, tháng và tuần trước khi in Báo cáo theo tuần!')

        if self.report_type == 'week':
            self.mau_in = 'template3'
        elif self.report_type == 'month':
            self.mau_in = 'template2'
        elif self.report_type == 'year':
            self.mau_in = 'template1'

        template_path = get_module_resource(
            'army_results_manager', 'static', 'src', 'word', f'{self.mau_in}.docx'
        )
        doc = Document(template_path)

        if self.report_type == 'week':
            self.replace_placeholder_with_text(doc, "{{week}}", self.week)
            self.replace_placeholder_with_text(doc, "{{month}}", self.month)

            # Lấy dữ liệu training days
            TrainingDay = self.env['training.day']
            domain = [
                ('year', '=', self.year),
                ('month_name', '=', f'Tháng {self.month}'),
                ('week_name', '=', f'Tuần {self.week}'),
                ('state', '=', 'approved'),
            ]

            records = TrainingDay.search(domain, order='day asc')

            if not records:
                raise UserError('Không tìm thấy dữ liệu!')

            table_index = 1
            if table_index >= len(doc.tables):
                raise UserError('Không tìm thấy table!')

            table = doc.tables[table_index]

            # Mapping weekday
            weekday_map = {
                '2': 'Hai',
                '3': 'Ba',
                '4': 'Tư',
                '5': 'Năm',
                '6': 'Sáu',
                '7': 'Bảy',
                'cn': 'Chủ nhật'
            }

            # NHÓM THEO COURSE_NAME VÀ NGÀY
            grouped_records = {}

            for record in records:
                weekday_text = weekday_map.get(record.weekday, record.weekday)
                day_str = record.day.strftime("%d/%m/%Y")
                key = (weekday_text, day_str)

                # Khởi tạo cấu trúc cho key nếu chưa tồn tại
                if key not in grouped_records:
                    grouped_records[key] = {}

                # Nhóm theo course_name
                course_name = record.course_name or "Không có tên khóa"
                if course_name not in grouped_records[key]:
                    grouped_records[key][course_name] = {
                        'lessons': [],  # Danh sách bài học
                        'total_hours': 0,  # Tổng số giờ
                        'times': []  # Danh sách thời gian
                    }

                # Thêm bài học nếu chưa có
                if record.lesson_name and record.lesson_name not in grouped_records[key][course_name]['lessons']:
                    grouped_records[key][course_name]['lessons'].append(record.lesson_name)

                # Cộng dồn tổng giờ
                grouped_records[key][course_name]['total_hours'] += (record.total_hours or 0)

                # Thêm thời gian
                for time_rec in record.time_ids:
                    if time_rec.start_time and time_rec.end_time:
                        # Chuyển đổi trực tiếp
                        start_h = int(time_rec.start_time)
                        start_m = int((time_rec.start_time - start_h) * 60)
                        end_h = int(time_rec.end_time)
                        end_m = int((time_rec.end_time - end_h) * 60)

                        time_str = f"{start_h:02d}:{start_m:02d} - {end_h:02d}:{end_m:02d}"
                        if time_str not in grouped_records[key][course_name]['times']:
                            grouped_records[key][course_name]['times'].append(time_str)

            # Điền vào bảng - CHỈ 1 HÀNG CHO MỖI NGÀY
            for (weekday, day_str), courses_data in grouped_records.items():
                # Thêm 1 hàng mới cho mỗi ngày
                new_row = table.add_row()

                # Điền weekday và ngày vào cùng 1 cell
                new_row.cells[0].text = f"{weekday}\n{day_str}"
                new_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Điền course_name và lessons vào cell[1]
                cell_content = new_row.cells[1]
                cell_content.text = ""

                # Điền hours vào cell[2]
                cell_hours = new_row.cells[2]
                cell_hours.text = ""

                # Điền time vào cell[3]
                cell_time = new_row.cells[3]
                cell_time.text = ""

                for course_name, course_data in courses_data.items():
                    # Thêm course_name với dấu :
                    p_course = cell_content.add_paragraph()
                    p_course.text = f"{course_name}:"

                    # Thêm tất cả lessons với dấu +
                    for lesson in course_data['lessons']:
                        p_lesson = cell_content.add_paragraph()
                        p_lesson.text = f"  + {lesson}"

                    # Thêm tổng hours cho course này
                    p_hour = cell_hours.add_paragraph()
                    p_hour.text = f"{course_data['total_hours']:g}" if course_data['total_hours'] else "0"
                    p_hour.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # Thêm times cho course này
                    for time_str in course_data['times']:
                        p_time = cell_time.add_paragraph()
                        p_time.text = time_str
                        p_time.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif self.report_type == 'month':
            self.replace_placeholder_with_text(doc, "{{year}}", self.year)
            self.replace_placeholder_with_text(doc, "{{month}}", self.month)

            # --- HÀM TẠO KHUNG 21 CỘT ---
            def make_time_slots():
                return {
                    1: [None] * 5,  # Tuần 1: T2-T6
                    2: [None] * 5,  # Tuần 2
                    3: [None] * 5,  # Tuần 3
                    4: [None] * 5,  # Tuần 4
                    5: [None] * 1  # Tuần 5: chỉ T2
                }

            def get_lower_letter(index):
                """Chuyển index thành a,b,c,..."""
                result = ""
                while index >= 0:
                    result = chr(index % 26 + 97) + result
                    index = index // 26 - 1
                return result

            def format_hours(hours):
                if hours == 0:
                    return ""

                # Làm tròn đến 1 chữ số thập phân
                rounded = round(hours, 1)

                # Nếu là số nguyên (ví dụ 3.0), bỏ phần thập phân
                if rounded == int(rounded):
                    return str(int(rounded))

                return str(rounded)

            TrainingDay = self.env['training.day']
            domain = [
                ('year', '=', self.year),
                ('month_name', '=', f'Tháng {self.month}'),
                ('state', '=', 'approved'),
            ]

            records = TrainingDay.search(domain)
            if not records:
                raise UserError('Không tìm thấy dữ liệu!')

            subject_columns = [
                "Chính trị",
                "Giáo dục pháp luật",
                "Hậu cần",
                "Kỹ thuật",
                "Điều lệnh",
                "Kỹ thuật CĐBB",
                "Bắn súng",
                "Thể lực chuyên môn",
                "Thể lực chung",
            ]

            # --- 1. Group dữ liệu theo plan -> subject ---
            grouped_data_table_0 = {}

            for rec in records:
                plan = rec.plan_name or "Không xác định"
                subject = rec.subject_name or "Không xác định"

                if plan not in grouped_data_table_0:
                    grouped_data_table_0[plan] = {}

                if subject not in grouped_data_table_0[plan]:
                    grouped_data_table_0[plan][subject] = {
                        "records": [],
                        "total_hours": 0,
                    }

                grouped_data_table_0[plan][subject]["records"].append(rec)
                grouped_data_table_0[plan][subject]["total_hours"] += rec.total_hours or 0

            # --- 2. Tạo dữ liệu cho table 0 ---
            table_0_data = []

            # Tạo table_0_data
            for idx, (plan_name, subjects) in enumerate(grouped_data_table_0.items()):
                row = [get_lower_letter(idx), plan_name]

                # Tính tổng số giờ và format
                total_hours = sum(subj_data["total_hours"] for subj_data in subjects.values())
                row.append(format_hours(total_hours))

                # Thêm số giờ cho từng môn (đã format)
                for subject in subject_columns:
                    hours = subjects.get(subject, {}).get("total_hours", 0)
                    row.append(format_hours(hours))

                # Thêm cột "Ghi chú" (để trống)
                row.append("")

                table_0_data.append(row)

            # --- 3. Ghi dữ liệu vào Word table 0 ---
            target_table = doc.tables[0]  # Bảng đầu tiên (index 0)

            # Bắt đầu ghi từ dòng 2 (sau 2 dòng header)
            start_row_index = 2

            for data_idx, data_row in enumerate(table_0_data):
                current_row_index = start_row_index + data_idx

                # Nếu cần thêm hàng mới (khi hết hàng trong table)
                while current_row_index >= len(target_table.rows):
                    target_table.add_row()

                # Lấy hàng tương ứng
                word_row = target_table.rows[current_row_index]

                # Ghi dữ liệu vào từng ô
                for col_idx, value in enumerate(data_row):
                    if col_idx < len(word_row.cells):
                        cell = word_row.cells[col_idx]
                        cell.text = str(value) if value else ""

                        # Căn giữa cho cột STT (cột 0) và các cột thời gian (cột 2-11)
                        if col_idx == 0 or (2 <= col_idx <= 11):
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # self.print_table(doc, 1)

            # --- MAP CHUYỂN CHUỖI THỨ → SỐ ---
            # --- KHUNG 21 CỘT ---
            def make_time_slots():
                return {
                    1: [None] * 5,  # Tuần 1: T2-T6
                    2: [None] * 5,
                    3: [None] * 5,
                    4: [None] * 5,
                    5: [None] * 1,  # Tuần 5: chỉ T2
                }

            weekday_map = {
                "thứ hai": 1,
                "thứ ba": 2,
                "thứ tư": 3,
                "thứ năm": 4,
                "thứ sáu": 5,
            }

            # --- BƯỚC 1: GOM NHÓM THEO MÔN HỌC VÀ BÀI HỌC ---
            grouped_data_table_1 = {
                "common_training": {},
                "private_training": {}
            }

            # --- Nhóm dữ liệu ---
            for rec in records:
                type_training = rec.type_training or "common_training"
                course = rec.course_name or "Không xác định"
                lesson_name = rec.lesson_name or "Không xác định"
                lesson_metadata_source = rec.course_id

                # Nhóm theo type_training
                if course not in grouped_data_table_1[type_training]:
                    grouped_data_table_1[type_training][course] = {
                        "lessons": {},
                        "course_total_hours": 0,
                        "plan_name": rec.plan_name or ""
                    }

                grouped_data_table_1[type_training][course]["course_total_hours"] += rec.total_hours or 0

                # Nhóm bài học
                lessons = grouped_data_table_1[type_training][course]["lessons"]
                if lesson_name not in lessons:
                    lessons[lesson_name] = {
                        "total_hours": 0,
                        "participant": lesson_metadata_source.participant_category_id.name or '',
                        "responsible": lesson_metadata_source.responsible_level_id.name or '',
                        "measure": lesson_metadata_source.measure or '',
                        "time_data_records": [],
                    }

                lesson_group = lessons[lesson_name]
                lesson_group["total_hours"] += rec.total_hours or 0
                lesson_group["time_data_records"].append({
                    "week": rec.week,
                    "weekday": rec.weekday,  # string kiểu "thứ hai", "thứ ba"
                    "hours": rec.total_hours or 0
                })

            # --- Hàm thêm môn học (đánh số linh hoạt) ---
            def add_course_rows(course_dict, start_course_idx=1, start_lesson_idx=None, is_private=False):
                rows = []
                course_idx = start_course_idx

                for course_name, course_data in sorted(course_dict.items(), key=lambda x: x[0]):
                    # Nếu là private, TT là số (1,2,..)
                    tt_course = str(course_idx) if is_private else str(course_idx)
                    course_row = {
                        'TT': tt_course,
                        'Nội dung huấn luyện': course_name.upper(),
                        'Thành phần tham gia': '',
                        'Cấp phụ trách': '',
                        'Tổng số (giờ)': format_hours(course_data["course_total_hours"]),
                        'Thời gian': [""] * 21,
                        'Biện pháp tiến hành': ''
                    }
                    rows.append(course_row)

                    # Các bài học (a,b,c) với khung số liệu 1.1, 1.2,...
                    lessons_list = sorted(course_data["lessons"].items(), key=lambda x: x[0])
                    for lesson_idx, (lesson_name, data) in enumerate(lessons_list, start=1):
                        # Tạo slot 21 cột
                        slots = make_time_slots()
                        for item in data["time_data_records"]:
                            week = int(item["week"]) if item["week"] else 0
                            weekday_str = (item["weekday"] or "").strip().lower()
                            weekday = weekday_map.get(weekday_str, 0)
                            hours = item["hours"]

                            if week in slots:
                                if week == 5:
                                    if weekday == 1:
                                        slots[5][0] = format_hours(hours)
                                else:
                                    if 1 <= weekday <= 5:
                                        slots[week][weekday - 1] = format_hours(hours)

                        flat_time_list = slots[1] + slots[2] + slots[3] + slots[4] + slots[5]
                        flat_time_list = [v if v else "" for v in flat_time_list]

                        lesson_tt = f"{course_idx}.{lesson_idx}" if is_private else get_lower_letter(lesson_idx - 1)
                        lesson_row = {
                            'TT': lesson_tt,
                            'Nội dung huấn luyện': lesson_name,
                            'Thành phần tham gia': data['participant'],
                            'Cấp phụ trách': data['responsible'],
                            'Tổng số (giờ)': format_hours(data['total_hours']),
                            'Thời gian': flat_time_list,
                            'Biện pháp tiến hành': data['measure']
                        }
                        rows.append(lesson_row)

                    course_idx += 1

                return rows, course_idx

            # --- Bước 2: tạo dữ liệu cho Word ---
            all_rows_to_write = []
            course_idx = 1

            # Thêm môn huấn luyện chung
            rows, course_idx = add_course_rows(grouped_data_table_1["common_training"], start_course_idx=course_idx)
            all_rows_to_write += rows

            # Thêm môn huấn luyện riêng
            if grouped_data_table_1["private_training"]:
                # Dòng B ở cột TT, HUẤN LUYỆN RIÊNG CÁC ĐỐI TƯỢNG ở nội dung huấn luyện, in đậm
                all_rows_to_write.append({
                    'TT': 'B',
                    'Nội dung huấn luyện': 'HUẤN LUYỆN RIÊNG CÁC ĐỐI TƯỢNG',
                    'Thành phần tham gia': '',
                    'Cấp phụ trách': '',
                    'Tổng số (giờ)': '',
                    'Thời gian': [""] * 21,
                    'Biện pháp tiến hành': ''
                })

                # Đánh số khóa huấn luyện bắt đầu từ 1
                rows, _ = add_course_rows(grouped_data_table_1["private_training"], start_course_idx=1, is_private=True)
                all_rows_to_write += rows

            # --- Bước 3: ghi vào Word ---
            target_table = doc.tables[1]
            start_row_index = 4
            current_row_index = start_row_index - 1

            for data_row in all_rows_to_write:
                current_row_index += 1
                while current_row_index >= len(target_table.rows):
                    target_table.add_row()
                word_row = target_table.rows[current_row_index]

                values = [
                             data_row['TT'],
                             data_row['Nội dung huấn luyện'],
                             data_row['Thành phần tham gia'],
                             data_row['Cấp phụ trách'],
                             data_row['Tổng số (giờ)'],
                         ] + data_row['Thời gian'] + [
                             data_row['Biện pháp tiến hành']
                         ]

                for col_idx, value in enumerate(values):
                    if col_idx < len(word_row.cells):
                        cell = word_row.cells[col_idx]
                        cell.text = str(value) if value else ""

                        if data_row['TT'] == 'B' and col_idx == 1:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True

                        # Căn giữa cho cột TT (0), Tổng số giờ (4), và các cột Thời gian (5-25)
                        if col_idx == 0 or col_idx == 4 or (5 <= col_idx <= 25):
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # table_1_data đã sẵn sàng để ghi vào doc.tables[1]

        elif self.report_type == 'year':
            def add_column(table):
                for row in table.rows:
                    last_cell = row.cells[-1]._tc
                    new_cell = deepcopy(last_cell)
                    row._tr.append(new_cell)

                    # CLEAR TEXT TRONG CELL COPY
                    cell_obj = row.cells[-1]
                    for p in cell_obj.paragraphs:
                        for run in p.runs:
                            run.text = ""
                        p.text = ""

            def merge_header_time(table, from_col, to_col):
                row = table.rows[0]
                start_cell = row.cells[from_col]
                for col in range(from_col + 1, to_col + 1):
                    start_cell.merge(row.cells[col])

            def center_cell(cell):
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            def fmt(dt):
                if not dt:
                    return ""
                return dt.strftime("%d/%m/%Y")

            self.replace_placeholder_with_text(doc, "{{year}}", self.year)

            TrainingDay = self.env['training.day']
            domain = [('year', '=', self.year), ('state', '=', 'approved')]
            records = TrainingDay.search(domain)

            if not records:
                raise UserError('Không tìm thấy dữ liệu!')

            self.print_table(doc, 0)

            grouped = {}
            for rec in records:
                if rec.plan_name not in grouped:
                    grouped[rec.plan_name] = rec
                else:
                    grouped[rec.plan_name] |= rec

            table = doc.tables[0]

            # =============================
            # 1. ADD COLUMNS IF NEEDED
            default_time_cols = 2
            needed_time_cols = len(grouped)

            # Số cột thời gian cần thêm
            extra_cols = needed_time_cols - default_time_cols
            if extra_cols > 0:
                for _ in range(extra_cols):
                    add_column(table)

            merge_header_time(table, 2, 1 + needed_time_cols)

            # =============================
            # 2. Fill plan_name vào dòng 1
            # =============================
            col_index = 2
            for plan_name in grouped.keys():
                cell = table.rows[1].cells[col_index]
                cell.text = plan_name
                center_cell(cell)
                col_index += 1

            # =============================
            # 3. Fill data
            # =============================

            col_index = 2
            for plan_name, items in grouped.items():
                start_date = min(items.mapped('plan_id.start_date'))
                end_date = max(items.mapped('plan_id.end_date'))

                # ----- Row 2: Bắt đầu -----
                cell = table.rows[2].cells[col_index]
                cell.text = fmt(start_date)
                center_cell(cell)

                # ----- Row 3: Kết thúc -----
                cell = table.rows[3].cells[col_index]
                cell.text = fmt(end_date)
                center_cell(cell)

                # ----- Row 4: Tổng số thời gian (số ngày) -----
                total_days = (end_date - start_date).days + 1
                cell = table.rows[4].cells[col_index]
                cell.text = str(total_days) + " ngày"
                center_cell(cell)

                col_index += 1

            # grouped = {}
            #
            # for rec in records:
            #     if rec.plan_name not in grouped:
            #         grouped[rec.plan_name] = rec
            #     else:
            #         grouped[rec.plan_name] |= rec
            #
            # table = doc.tables[0]
            #
            # # Add columns if needed
            # required_cols = 2 + len(grouped)
            # current_cols = len(table.rows[0].cells)
            # if required_cols > current_cols:
            #     diff = required_cols - current_cols
            #     for _ in range(diff):
            #         add_column(table)
            #
            # # Fill plan_name into row[1]
            # col_index = 2
            # for plan_name in grouped.keys():
            #     table.rows[1].cells[col_index].text = plan_name
            #     col_index += 1
            #
            # # Fill data for each plan_name
            # col_index = 2
            # for plan_name, items in grouped.items():
            #     start_date = min(items.mapped('plan_id.start_date'))
            #     end_date = max(items.mapped('plan_id.end_date'))
            #     # total_days = sum(items.mapped('total_days'))
            #     # weeks = total_days // 7
            #     # holidays = sum(items.mapped('holidays'))
            #
            #     table.rows[2].cells[col_index].text = str(start_date)
            #     table.rows[3].cells[col_index].text = str(end_date)
            #     # table.rows[4].cells[col_index].text = str(total_days)
            #     # table.rows[5].cells[col_index].text = str(weeks)
            #     # table.rows[6].cells[col_index].text = str(total_days)
            #     # table.rows[7].cells[col_index].text = str(holidays)
            #
            #     col_index += 1

            # Lấy set của tất cả plan_id (unique plans)

        file_data = BytesIO()
        doc.save(file_data)
        file_data.seek(0)
        data = base64.b64encode(file_data.read())

        if hasattr(self, 'week') and self.week:
            # Có tuần: Báo cáo huấn luyện tuần X tháng Y năm Z
            report_name = f'Bao_cao_huan_luyen_tuan_{self.week}_thang_{self.month}_nam_{self.year}.docx'
        elif hasattr(self, 'month') and self.month:
            # Có tháng: Báo cáo huấn luyện tháng X năm Y
            report_name = f'Bao_cao_huan_luyen_thang_{self.month}_nam_{self.year}.docx'
        else:
            # Chỉ có năm: Báo cáo huấn luyện năm X
            report_name = f'Bao_cao_huan_luyen_nam_{self.year}.docx'

        attachment = self.env['ir.attachment'].create({
            'name': report_name,
            'type': 'binary',
            'datas': data,
            'res_model': self._name,
            'res_id': self.id,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        })

        return {
            'type': 'ir.actions.act_url',
            'url': f'/web/content/{attachment.id}?download=true',
            'target': 'new',
        }
